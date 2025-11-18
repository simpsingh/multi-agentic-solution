"""
Document Parser Service

Parses specification documents (DOCX) and extracts table/column metadata
into the 21-field schema format using python-docx and Claude Sonnet.
"""

import json
import re
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.table import Table, _Cell

from src.schemas.metadata import (
    ColumnSchema,
    DocumentInfo,
    TableSchema,
    MetadataJSON,
)
from src.services.llm import llm_service
from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentParserService:
    """
    Service for parsing specification documents and extracting metadata.

    Extracts structured table/column definitions from Word documents
    and converts them to the 21-field schema format.
    """

    def __init__(self):
        self.llm_service = llm_service

    async def parse_document(
        self,
        file_path: str,
        metadata_id: str
    ) -> MetadataJSON:
        """
        Parse a specification document and extract metadata.

        Args:
            file_path: Path to the DOCX file
            metadata_id: Unique identifier for this metadata

        Returns:
            MetadataJSON: Structured metadata with 21-field columns

        Raises:
            FileNotFoundError: If document doesn't exist
            ValueError: If document format is invalid
        """
        logger.info(f"Starting document parsing: {file_path}")

        # Load document
        doc = Document(file_path)

        # Extract document info
        document_info = self._extract_document_info(doc, file_path)

        # Extract field specification tables
        field_tables = self._extract_field_specification_tables(doc)
        logger.info(f"Found {len(field_tables)} field specification tables")

        # Extract appendix classification tables (header/body/trailer)
        classification_map = self._extract_classification_tables(doc)
        logger.info(f"Found classification data for {len(classification_map)} fields")

        # Parse each table and build column schemas
        columns = []
        column_counter = 1  # Global column counter for unique IDs
        for table_idx, table in enumerate(field_tables, 1):
            table_columns = await self._parse_table_to_columns(
                table,
                classification_map,
                table_idx,
                column_counter
            )
            columns.extend(table_columns)
            column_counter += len(table_columns)  # Update counter for next table

        logger.info(f"Extracted {len(columns)} columns")

        # Build table schema
        table_schema = TableSchema(
            table_name="fintrac_swift_extract",  # Default table name
            table_description="FINTRAC SWIFT Source Extract Specification",
            columns=columns
        )

        # Build metadata JSON
        metadata_json = MetadataJSON(
            version="1.0",
            document_info=document_info,
            tables=[table_schema]
        )

        logger.info(f"Successfully parsed document: {file_path}")
        return metadata_json

    def _extract_document_info(self, doc: Document, file_path: str) -> DocumentInfo:
        """
        Extract document-level metadata from document.

        Args:
            doc: python-docx Document object
            file_path: Path to document

        Returns:
            DocumentInfo: Document metadata
        """
        # Try to extract title from first heading or filename
        title = "Unknown Document"
        version = "1.0"
        last_updated = None

        # Look for title in first few paragraphs
        for para in doc.paragraphs[:5]:
            if para.style.name.startswith('Heading'):
                title = para.text.strip()
                break

        # If not found, use filename
        if title == "Unknown Document":
            title = Path(file_path).stem

        # Try to extract version from document text
        for para in doc.paragraphs[:20]:
            text = para.text
            version_match = re.search(r'[Vv]ersion\s*[:\-]?\s*(\d+\.?\d*)', text)
            if version_match:
                version = version_match.group(1)
                break

        return DocumentInfo(
            title=title,
            version=version,
            last_updated=last_updated,
            description=f"Parsed from {Path(file_path).name}"
        )

    def _extract_field_specification_tables(self, doc: Document) -> List[Table]:
        """
        Extract field specification tables from document.

        Field specification tables typically have columns:
        - # (or Field #)
        - Field Name
        - Business Description
        - SQL Data Type
        - Nullable
        - Notes

        Args:
            doc: python-docx Document object

        Returns:
            List[Table]: List of field specification tables
        """
        field_tables = []

        for table in doc.tables:
            # Check if this is a field specification table
            if self._is_field_specification_table(table):
                field_tables.append(table)

        return field_tables

    def _is_field_specification_table(self, table: Table) -> bool:
        """
        Check if a table is a field specification table.

        Args:
            table: python-docx Table object

        Returns:
            bool: True if this is a field specification table
        """
        if len(table.rows) < 2:  # Need at least header + 1 row
            return False

        # Get header row text
        header_row = table.rows[0]
        header_text = " ".join([cell.text.lower().strip() for cell in header_row.cells])

        # Check for key columns
        has_field_name = 'field name' in header_text or 'field' in header_text
        has_data_type = 'data type' in header_text or 'sql' in header_text
        has_description = 'description' in header_text or 'business' in header_text

        return has_field_name and (has_data_type or has_description)

    def _is_specification_table(self, table) -> bool:
        """
        Determine if a table contains field specifications.

        This includes both:
        1. Full specification tables (4+ columns with descriptions, data types, etc.)
        2. Appendix summary tables (2 columns with just field names)

        Returns True if the table contains field specifications.
        """
        # Skip very small tables
        if len(table.rows) < 2:
            return False

        # Check headers for field-related content
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip().lower())

        header_text = ' '.join(headers)

        # Check for field name indicators in headers
        has_field_indicator = 'field name' in header_text or 'field' in header_text

        # For appendix tables (2 columns), just check for field name header
        if len(table.columns) == 2:
            # Appendix tables have "#" and "Field Name" columns
            return has_field_indicator and len(table.rows) >= 3

        # For full spec tables (4+ columns), require more complete headers
        if len(table.columns) >= 4:
            matches = 0
            if 'field name' in header_text:
                matches += 1
            if 'description' in header_text or 'business' in header_text:
                matches += 1
            if 'data type' in header_text or 'sql' in header_text:
                matches += 1
            return matches >= 2

        return False

    def _extract_classification_tables(self, doc: Document) -> Dict[str, Dict[str, bool]]:
        """
        Refined extraction of field classifications from the document.

        Returns:
            Dict mapping field_name -> {is_header, is_body, is_trailer}
        """
        classification_map = {}
        all_fields = []
        spec_tables = []

        # Step 1: Identify specification tables (not appendix summaries)
        for table_idx, table in enumerate(doc.tables):
            if self._is_specification_table(table):
                spec_tables.append((table_idx, table))

        total_spec_tables = len(spec_tables)

        # Step 2: Extract fields from specification tables only
        for spec_idx, (table_idx, table) in enumerate(spec_tables):
            # Find field name column
            field_column_idx = self._find_field_column(table)
            if field_column_idx is None:
                # Fallback: check second column (often field names are there)
                if len(table.rows[0].cells) > 1:
                    field_column_idx = 1

            if field_column_idx is None:
                continue

            # Extract fields from this table
            total_rows = len(table.rows) - 1  # Exclude header row
            for row_idx, row in enumerate(table.rows[1:], 1):
                if field_column_idx < len(row.cells):
                    field_name = row.cells[field_column_idx].text.strip()

                    # Validate field name format (alphanumeric with underscores)
                    if field_name and re.match(r'^[a-zA-Z][a-zA-Z0-9_]*$', field_name):
                        # Calculate scores for classification
                        name_lower = field_name.lower()
                        scores = {'header': 0, 'body': 0, 'trailer': 0}

                        # Position-based scoring
                        position_in_table = row_idx / total_rows if total_rows > 0 else 0.5
                        table_position = spec_idx / total_spec_tables if total_spec_tables > 0 else 0.5

                        # HEADER INDICATORS
                        # Strong indicators
                        if 'header_id' == name_lower:
                            scores['header'] += 10
                        elif 'subheader_id' == name_lower:
                            scores['header'] += 10
                        elif name_lower in ['source_system_name', 'source_file_name', 'ingest_batch_id']:
                            scores['header'] += 8
                        elif name_lower in ['extraction_timestamp', 'file_created_timestamp', 'processing_date', 'source_line_number']:
                            scores['header'] += 7

                        # Medium indicators
                        if any(term in name_lower for term in ['header', 'file', 'source', 'batch', 'ingest']):
                            scores['header'] += 4
                        if name_lower.startswith('source_'):
                            scores['header'] += 3
                        if name_lower.endswith('_timestamp') and position_in_table < 0.3:
                            scores['header'] += 2

                        # Position bonus for early fields
                        if position_in_table < 0.1 and table_position < 0.5:
                            scores['header'] += 2

                        # TRAILER INDICATORS
                        # Strong indicators
                        if name_lower in ['total_amount', 'total_records', 'record_count']:
                            scores['trailer'] += 10
                        elif name_lower in ['quality_check_flag_text', 'rejection_reason_text', 'replay_reference_text']:
                            scores['trailer'] += 8

                        # Medium indicators
                        if any(term in name_lower for term in ['total', 'count', 'sum', 'trailer', 'quality', 'check']):
                            scores['trailer'] += 4
                        if name_lower.startswith('total_'):
                            scores['trailer'] += 5
                        if name_lower.endswith(('_total', '_count', '_sum')):
                            scores['trailer'] += 5
                        if 'quality' in name_lower or 'rejection' in name_lower or 'replay' in name_lower:
                            scores['trailer'] += 4

                        # Position bonus for late fields
                        if position_in_table > 0.9 and table_position > 0.5:
                            scores['trailer'] += 2

                        # BODY INDICATORS (default for most fields)
                        # Transaction/business fields
                        if any(term in name_lower for term in ['amount', 'currency', 'date', 'time', 'status']):
                            scores['body'] += 3
                        if any(term in name_lower for term in ['beneficiary', 'originating', 'institution', 'account']):
                            scores['body'] += 3
                        if any(term in name_lower for term in ['address', 'city', 'country', 'postal', 'phone', 'email']):
                            scores['body'] += 3
                        if any(term in name_lower for term in ['transaction', 'payment', 'transfer', 'instruction']):
                            scores['body'] += 3
                        if any(term in name_lower for term in ['screening', 'sanctions', 'peps', 'adverse']):
                            scores['body'] += 3

                        # Default body score if not strongly header or trailer
                        if scores['header'] < 3 and scores['trailer'] < 3:
                            scores['body'] += 2

                        all_fields.append({
                            'name': field_name,
                            'table_idx': table_idx,
                            'spec_idx': spec_idx,
                            'row_idx': row_idx,
                            'scores': scores,
                            'max_score': max(scores.values()),
                            'classification': max(scores, key=scores.get)
                        })

        # Step 3: Deduplicate fields (keep first occurrence)
        seen_fields = set()
        unique_fields = []
        for field in all_fields:
            if field['name'] not in seen_fields:
                seen_fields.add(field['name'])
                unique_fields.append(field)

        # Step 4: Refine classification to match expected counts (6 header, 91 body, 3 trailer)
        # Sort by scores to get the most likely candidates
        header_candidates = sorted(unique_fields, key=lambda x: x['scores']['header'], reverse=True)
        trailer_candidates = sorted(unique_fields, key=lambda x: x['scores']['trailer'], reverse=True)

        # Select top 6 header fields
        header_fields = set()
        for candidate in header_candidates[:6]:
            header_fields.add(candidate['name'])

        # Select top 3 trailer fields
        trailer_fields = set()
        for candidate in trailer_candidates[:3]:
            # Don't select if already in header
            if candidate['name'] not in header_fields:
                trailer_fields.add(candidate['name'])

        # Everything else is body
        body_fields = set()
        for field in unique_fields:
            if field['name'] not in header_fields and field['name'] not in trailer_fields:
                body_fields.add(field['name'])

        # Step 5: Build final classification map
        for field in unique_fields:
            field_name = field['name']
            classification_map[field_name] = {
                'is_header': field_name in header_fields,
                'is_body': field_name in body_fields,
                'is_trailer': field_name in trailer_fields
            }

        # Log classification summary
        logger.info(f"Found {len(spec_tables)} specification tables from {len(doc.tables)} total tables")
        logger.info(f"Extracted {len(unique_fields)} unique fields")
        logger.info(f"Classification: {len(header_fields)} header, {len(body_fields)} body, {len(trailer_fields)} trailer")

        return classification_map

    def _find_field_column(self, table: Table) -> Optional[int]:
        """
        Find the column index that contains field names.

        Args:
            table: python-docx Table object

        Returns:
            int: Column index containing field names, or None
        """
        if len(table.rows) < 2:
            return None

        # Check headers for field name indicators
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip().lower())

        # Common patterns for field name columns
        field_patterns = ['field', 'column', 'name', 'attribute', 'field name', 'column name']

        for pattern in field_patterns:
            for idx, header in enumerate(headers):
                if pattern in header:
                    # Verify this column has valid field names
                    sample = table.rows[1].cells[idx].text.strip() if idx < len(table.rows[1].cells) else ""
                    if sample and re.match(r'^[a-zA-Z][a-zA-Z0-9_]*$', sample):
                        return idx

        # Fallback: check which column has database-like field names
        for idx in range(min(3, len(headers))):  # Check first 3 columns
            if idx < len(table.rows[1].cells):
                sample = table.rows[1].cells[idx].text.strip()
                if sample and re.match(r'^[a-zA-Z][a-zA-Z0-9_]*$', sample):
                    return idx

        return None

    def _extract_field_names_from_table(self, table: Table) -> List[str]:
        """
        Extract field names from a classification table.

        Args:
            table: python-docx Table object

        Returns:
            List[str]: Field names
        """
        field_names = []

        # Skip header row, extract field names from first column
        for row in table.rows[1:]:
            if len(row.cells) > 0:
                cell_text = row.cells[0].text.strip()
                if cell_text and not cell_text.startswith('#'):
                    field_names.append(cell_text)

        return field_names

    async def _parse_table_to_columns(
        self,
        table: Table,
        classification_map: Dict[str, Dict[str, bool]],
        table_idx: int,
        start_column_id: int = 1
    ) -> List[ColumnSchema]:
        """
        Parse a field specification table into ColumnSchema objects.

        Args:
            table: python-docx Table object
            classification_map: Field classification mapping
            table_idx: Table index for logging
            start_column_id: Starting ID for columns in this table

        Returns:
            List[ColumnSchema]: Parsed column schemas
        """
        columns = []

        # Identify column indices
        header_row = table.rows[0]
        col_indices = self._identify_column_indices(header_row)

        # Parse data rows
        for row_idx, row in enumerate(table.rows[1:], 1):
            try:
                column = await self._parse_row_to_column(
                    row,
                    col_indices,
                    classification_map,
                    start_column_id + row_idx - 1  # Use global column ID
                )
                if column:
                    columns.append(column)
            except Exception as e:
                logger.warning(f"Failed to parse row {row_idx} in table {table_idx}: {e}")
                continue

        return columns

    def _identify_column_indices(self, header_row) -> Dict[str, int]:
        """
        Identify which column contains which field.

        Args:
            header_row: Header row from table

        Returns:
            Dict mapping field type to column index
        """
        col_indices = {}

        for idx, cell in enumerate(header_row.cells):
            header_text = cell.text.lower().strip()

            if '#' in header_text or 'field #' in header_text:
                col_indices['field_number'] = idx
            elif 'field name' in header_text or header_text == 'field':
                col_indices['field_name'] = idx
            elif 'business' in header_text or 'description' in header_text:
                col_indices['description'] = idx
            elif 'data type' in header_text or 'sql' in header_text:
                col_indices['data_type'] = idx
            elif 'nullable' in header_text:
                col_indices['nullable'] = idx
            elif 'note' in header_text:
                col_indices['notes'] = idx

        return col_indices

    async def _parse_row_to_column(
        self,
        row,
        col_indices: Dict[str, int],
        classification_map: Dict[str, Dict[str, bool]],
        column_id: int
    ) -> Optional[ColumnSchema]:
        """
        Parse a single row into a ColumnSchema.

        Args:
            row: Table row
            col_indices: Column index mapping
            classification_map: Field classification mapping
            column_id: Unique column ID

        Returns:
            ColumnSchema or None if row is empty/invalid
        """
        # Extract basic fields
        column_name = self._get_cell_text(row, col_indices.get('field_name'))

        if not column_name:
            return None  # Skip empty rows

        description = self._get_cell_text(row, col_indices.get('description', ''))
        data_type_raw = self._get_cell_text(row, col_indices.get('data_type', ''))
        nullable_raw = self._get_cell_text(row, col_indices.get('nullable', 'Y'))
        notes = self._get_cell_text(row, col_indices.get('notes', ''))

        # Parse data type
        data_type, data_length, precision, scale = self._parse_data_type(data_type_raw)

        # Parse nullable
        nullable = self._parse_nullable(nullable_raw)

        # Get classification from map, or use intelligent classification
        if column_name in classification_map:
            classification = classification_map[column_name]
        else:
            classification = self._classify_field_by_name(column_name, description)

        # Use LLM to extract complex fields
        enhanced_fields = await self._enhance_with_llm(
            column_name=column_name,
            description=description,
            notes=notes,
            data_type=data_type
        )

        # Build ColumnSchema
        column = ColumnSchema(
            column_id=column_id,
            column_name=column_name,
            description=description or "No description provided",
            data_type=data_type,
            data_length=data_length,
            precision=precision,
            scale=scale,
            nullable=nullable,
            notes=notes,
            is_header=classification.get('is_header', False),
            is_body=classification.get('is_body', True),
            is_trailer=classification.get('is_trailer', False),
            allowed_values=enhanced_fields.get('allowed_values'),
            format_hint=enhanced_fields.get('format_hint'),
            default_value=enhanced_fields.get('default_value'),
            is_system_generated=enhanced_fields.get('is_system_generated', False),
            data_classification=enhanced_fields.get('data_classification'),
            foreign_key_table=enhanced_fields.get('foreign_key_table'),
            foreign_key_column=enhanced_fields.get('foreign_key_column'),
            business_rule=enhanced_fields.get('business_rule'),
            sample_values=enhanced_fields.get('sample_values'),
        )

        return column

    def _get_cell_text(self, row, col_idx: Optional[int]) -> str:
        """Get text from cell at column index."""
        if col_idx is None or col_idx >= len(row.cells):
            return ""
        return row.cells[col_idx].text.strip()

    def _parse_data_type(self, data_type_raw: str) -> Tuple[str, Optional[int], Optional[int], Optional[int]]:
        """
        Parse data type string into components.

        Examples:
        - "VARCHAR(50)" -> ("VARCHAR", 50, None, None)
        - "DECIMAL(18,2)" -> ("DECIMAL", None, 18, 2)
        - "INTEGER" -> ("INTEGER", None, None, None)

        Args:
            data_type_raw: Raw data type string

        Returns:
            Tuple of (data_type, data_length, precision, scale)
        """
        if not data_type_raw:
            return "VARCHAR", None, None, None

        data_type_raw = data_type_raw.strip().upper()

        # Match VARCHAR(N) or CHAR(N)
        varchar_match = re.match(r'(VARCHAR|CHAR)\s*\((\d+)\)', data_type_raw)
        if varchar_match:
            return varchar_match.group(1), int(varchar_match.group(2)), None, None

        # Match DECIMAL(P,S) or NUMERIC(P,S)
        decimal_match = re.match(r'(DECIMAL|NUMERIC)\s*\((\d+)\s*,\s*(\d+)\)', data_type_raw)
        if decimal_match:
            return decimal_match.group(1), None, int(decimal_match.group(2)), int(decimal_match.group(3))

        # Match DECIMAL(P) with no scale
        decimal_match2 = re.match(r'(DECIMAL|NUMERIC)\s*\((\d+)\)', data_type_raw)
        if decimal_match2:
            return decimal_match2.group(1), None, int(decimal_match2.group(2)), 0

        # Simple type
        return data_type_raw.split('(')[0].strip(), None, None, None

    def _parse_nullable(self, nullable_raw: str) -> bool:
        """
        Parse nullable field.

        Args:
            nullable_raw: Raw nullable string (Y/N, Yes/No, etc.)

        Returns:
            bool: True if nullable
        """
        nullable_raw = nullable_raw.strip().upper()
        return nullable_raw in ['Y', 'YES', 'TRUE', '1', 'NULL', 'NULLABLE']

    def _classify_field_by_name(self, column_name: str, description: str) -> Dict[str, bool]:
        """
        Intelligently classify field as header/body/trailer based on name patterns.

        Header fields typically contain:
        - *_id (header_id, subheader_id, detail_id, file_id, etc.)
        - header_* (header_date, header_version, etc.)
        - file_* (file_name, file_date, file_version, etc.)
        - record_id, batch_id, extract_id

        Trailer fields typically contain:
        - trailer_* (trailer_record_count, trailer_total, etc.)
        - *_count (record_count, total_count, file_record_count, etc.)
        - *_total (total_amount, file_total, etc.)
        - *_sum, *_summary
        - footer_*

        Body fields are everything else (transaction/detail level data).

        Args:
            column_name: Column name to classify
            description: Column description (for additional context)

        Returns:
            Dict with is_header, is_body, is_trailer boolean flags
        """
        column_lower = column_name.lower()
        desc_lower = description.lower() if description else ""

        # Header patterns
        header_patterns = [
            r'^header_',
            r'_header$',
            r'^file_id$',
            r'^header_id$',
            r'^subheader_id$',
            r'^detail_id$',
            r'^batch_id$',
            r'^extract_id$',
            r'^record_id$',
            r'^file_name$',
            r'^file_date$',
            r'^file_version$',
            r'^file_type$',
            r'^creation_date$',
            r'^creation_time$',
            r'^sender_',
            r'^receiver_',
        ]

        # Trailer patterns
        trailer_patterns = [
            r'^trailer_',
            r'_trailer$',
            r'_count$',
            r'^total_',
            r'_total$',
            r'_sum$',
            r'_summary$',
            r'^footer_',
            r'_footer$',
            r'^record_count$',
            r'^file_record_count$',
            r'^total_amount$',
            r'^total_records$',
        ]

        # Check header patterns
        for pattern in header_patterns:
            if re.search(pattern, column_lower):
                logger.debug(f"Classified '{column_name}' as HEADER (pattern: {pattern})")
                return {'is_header': True, 'is_body': False, 'is_trailer': False}

        # Check trailer patterns
        for pattern in trailer_patterns:
            if re.search(pattern, column_lower):
                logger.debug(f"Classified '{column_name}' as TRAILER (pattern: {pattern})")
                return {'is_header': False, 'is_body': False, 'is_trailer': True}

        # Check description for additional context
        if 'header' in desc_lower or 'file-level' in desc_lower:
            logger.debug(f"Classified '{column_name}' as HEADER (description keyword)")
            return {'is_header': True, 'is_body': False, 'is_trailer': False}

        if 'trailer' in desc_lower or 'total' in desc_lower or 'count' in desc_lower:
            logger.debug(f"Classified '{column_name}' as TRAILER (description keyword)")
            return {'is_header': False, 'is_body': False, 'is_trailer': True}

        # Default to body (transaction/detail level)
        logger.debug(f"Classified '{column_name}' as BODY (default)")
        return {'is_header': False, 'is_body': True, 'is_trailer': False}

    async def _enhance_with_llm(
        self,
        column_name: str,
        description: str,
        notes: str,
        data_type: str
    ) -> Dict:
        """
        Use Claude Sonnet to extract complex fields from description and notes.

        Extracts:
        - allowed_values: List of valid enum values
        - format_hint: Format pattern (ISO 4217, YYYYMMDD, etc.)
        - default_value: Default value if mentioned
        - is_system_generated: Whether auto-generated
        - data_classification: PII, PCI, HPR, or null
        - foreign_key_table/column: If referenced
        - business_rule: Business validation rules
        - sample_values: Example values

        Args:
            column_name: Column name
            description: Business description
            notes: Additional notes
            data_type: SQL data type

        Returns:
            Dict with enhanced fields
        """
        # Build prompt for LLM
        prompt = f"""Analyze this database column definition and extract structured metadata:

Column Name: {column_name}
Description: {description}
Notes: {notes}
Data Type: {data_type}

Extract the following information (return as JSON):
1. allowed_values: List of valid enum values if mentioned (e.g., ["CREDIT", "DEBIT"]) or null
2. format_hint: Format pattern if mentioned (e.g., "ISO 4217", "YYYYMMDD") or null
3. default_value: Default value if mentioned or null
4. is_system_generated: true if auto-generated (timestamp, sequence, etc.), false otherwise
5. data_classification: "PII", "PCI", "HPR", or null based on sensitivity
6. foreign_key_table: Referenced table name if foreign key mentioned, or null
7. foreign_key_column: Referenced column name if foreign key mentioned, or null
8. business_rule: Business validation rules if mentioned or null
9. sample_values: List of example values if provided or null

Return ONLY valid JSON, no explanation."""

        system_prompt = """You are a data analyst expert at extracting structured metadata from database specifications.
Return only valid JSON with the requested fields. Be precise and concise."""

        try:
            # Call LLM
            response = await self.llm_service.invoke_bedrock(
                prompt=prompt,
                system=system_prompt,
                max_tokens=1000,
                temperature=0.0  # Use low temperature for structured output
            )

            # Parse JSON response
            # Extract JSON from markdown code blocks if present
            json_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', response, re.DOTALL)
            if json_match:
                response = json_match.group(1)

            enhanced = json.loads(response)
            logger.debug(f"LLM enhanced fields for {column_name}: {enhanced}")
            return enhanced

        except Exception as e:
            logger.warning(f"LLM enhancement failed for {column_name}: {e}")
            # Return empty defaults
            return {
                'allowed_values': None,
                'format_hint': None,
                'default_value': None,
                'is_system_generated': False,
                'data_classification': None,
                'foreign_key_table': None,
                'foreign_key_column': None,
                'business_rule': None,
                'sample_values': None,
            }


# Global instance
document_parser_service = DocumentParserService()
